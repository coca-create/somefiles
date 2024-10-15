import os
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
import tempfile
import json
from pydub import AudioSegment
import gradio as gr
from faster_whisper import WhisperModel
import zipfile
from docx import Document
import re
from openpyxl.styles import Alignment, Font, PatternFill
import pandas as pd
from tqdm import tqdm
from datetime import datetime
import traceback

def get_audio_duration(filepath):
    try:
        print(filepath)
        audio = AudioSegment.from_file(filepath)
        duration = len(audio) / 1000.0
        return duration

    except Exception as e:
        print(f"get_audio_durationでerror:{e}")
        traceback.print_exec()
        return None

def format_timestamp(seconds):
    hrs, secs = divmod(seconds, 3600)
    mins, secs = divmod(secs, 60)
    millis = int((secs % 1) * 1000)
    return f"{int(hrs):02}:{int(mins):02}:{int(secs):02},{millis:03}"

#dataframe追加
def parse_srt_c(srt_content):
    pattern = r'(\d+)\n(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\n(.*?)\n\n'
    matches = re.findall(pattern, srt_content, re.DOTALL)
    
    subtitles = []
    for match in matches:
        subtitles.append({
            'ID': int(match[0]),
            'Start': match[1],
            'End': match[2],
            'Text': match[3].replace('\n', ' ')
        })
    
    return subtitles

def dataframe_to_html_table(df):
    return df.to_html(index=False)

# SRTファイルからExcelファイルを作成する関数
def create_excel_from_srt_c(srt_content, input_file_name):
    excel_file_name = f"{input_file_name}_srt.xlsx"
    english_subtitles = parse_srt_c(srt_content)

    data = []
    for eng in english_subtitles:
        data.append({
            'ID': eng['ID'],
            'Start': eng['Start'],
            'End': eng['End'],
            'English Subtitle': eng['Text']
        })

    df = pd.DataFrame(data)
    temp_dir = tempfile.gettempdir()
    excel_file_path = os.path.join(temp_dir, excel_file_name)
    
    with pd.ExcelWriter(excel_file_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Subtitles')
        workbook = writer.book
        worksheet = writer.sheets['Subtitles']

        column_widths = {'A': 7, 'B': 25, 'C': 25, 'D': 90, 'E': 90}
        for column, width in column_widths.items():
            worksheet.column_dimensions[column].width = width

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            for cell in row:
                if cell.column_letter == 'A':
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                elif cell.column_letter in ['B', 'C']:
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                elif cell.column_letter in ['D', 'E']:
                    cell.alignment = Alignment(horizontal='left', vertical='center')

        for row in worksheet.iter_rows(min_row=2, max_row=len(df) + 1):
            worksheet.row_dimensions[row[0].row].height = 30

        header_font = Font(bold=True)
        for cell in worksheet["1:1"]:
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
            cell.fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")

    
    return excel_file_path, df

'''def exe_for_gradio(srt_content, input_file_name='Noname'):
    excel_filepath, df_display = create_excel_from_srt_c(srt_content, input_file_name)
    return df_display'''

def transcribe(File, Model, Computing, Lang, BeamSize, VadFilter, device, progress=gr.Progress()):
    if not File:
        print("ファイルが提供されていません")
        return"", "", "", [], [], "", "", "", "", "",""

    print(File)
    FileName = File
    if Lang == "日本語":
        Lang = "ja"
    else:
        Lang = "en"
    
    model = WhisperModel(Model, device=device, compute_type=Computing)
    print(f"using:{device}")
    segments, _ = model.transcribe(File, word_timestamps=True, beam_size=BeamSize, initial_prompt="Hello, I am Scott.", language=Lang, vad_filter=VadFilter)

        #error_message = f"文字起こし中にエラーが発生しました: {e}"
        #return error_message, "", "", [], [], "", "", "", "", "",""

    total_duration = get_audio_duration(File)
    if isinstance(total_duration, str):  # get_audio_duration関数がエラーメッセージを返した場合
        return total_duration, "", "", [], [], "", "", "", "", "",""
    
    words_data = []

    try:
        # Initialize tqdm progress bar
        progress_bar = tqdm(total=total_duration, unit="s", position=0, leave=True, desc="処理進行状況")
        last_update_time = 0

        # 初期更新
        progress(0)
        progress_bar.update(0)

        for segment in segments:
            for word in segment.words:
                word_info = {
                    "start": word.start,
                    "end": word.end,
                    "word": word.word
                }
                words_data.append(word_info)
            
            # Update progress at a reasonable frequency
            segment_progress = segment.end - last_update_time
            if segment_progress >= 1.0:  # Update every 1 second
                progress_bar.update(segment_progress)
                progress(segment.end / total_duration)
                last_update_time = segment.end        
                progress_bar.set_postfix({"progress": segment.end / total_duration * 100})

        progress_bar.update(total_duration - last_update_time)
        progress(1.0)  # Ensure progress is 100% at the end
        # Close tqdm progress bar
        progress_bar.close()

    except Exception as e:
        print(f"進捗バー更新中にエラーが発生しました(ループ処理中のエラー): {e}")
        traceback.print_exc()
        return "", "", "", [], [], "", "", "", "", "",""
 
     

    try:
        #print("ファイル処理-Dot変換")
        for word_info in words_data:
            word_info["word"] = word_info["word"].replace(" Dr.", " Dr★").replace(" dr.", " dr★")

        # 前処理: words_data内の各wordの中から★を削除する
        cleaned_words_data = []
        for word_info in words_data:
            cleaned_word_info = {
                "start": word_info["start"],
                "end": word_info["end"],
                "word": word_info["word"].replace("★", ".")
            }
            cleaned_words_data.append(cleaned_word_info)

        input_file_name = os.path.splitext(os.path.basename(File))[0]
        
        timestamp_patch = datetime.now().strftime("%Y%m%d%H%M%S")
        temp_dir = os.path.join(tempfile.gettempdir(), f"tempdir_{timestamp_patch}")
        os.makedirs(temp_dir, exist_ok=True)
        
        json_output_file_name = f"{input_file_name}.json"
        json_output_path = os.path.join(temp_dir, json_output_file_name)
        # JSONファイルへの書き込み
        with open(json_output_path, 'w', encoding='utf-8') as f:
           json.dump(cleaned_words_data, f, ensure_ascii=False, indent=4)
           
        # 書き込んだJSONデータの表示（デバッグ用）
        json_content = json.dumps(cleaned_words_data, ensure_ascii=False, indent=4)

        srt_entries = []
        entry_number = 1
        segment_text = ""
        segment_start = None
        segment_end = None

        for word_info in words_data:
            if segment_start is None:
                segment_start = word_info["start"]
            
            segment_text += word_info["word"]
            segment_end = word_info["end"]
            
            if word_info["word"].endswith('.'):
                srt_entries.append({
                    "number": entry_number,
                    "start": segment_start,
                    "end": segment_end,
                    "text": segment_text.strip()
                })
                entry_number += 1
                segment_text = ""
                segment_start = None

        if segment_text.strip():
            srt_entries.append({
                "number": entry_number,
                "start": segment_start,
                "end": segment_end,
                "text": segment_text.strip()
            })

        srt_output_file_name = f"{input_file_name}.srt"
        srt_output_path = os.path.join(temp_dir, srt_output_file_name)

        with open(srt_output_path, 'w', encoding='utf-8') as f:
            for entry in srt_entries:
                start_time = format_timestamp(entry["start"])
                end_time = format_timestamp(entry["end"])
                text = entry['text'].replace(" Dr★", " Dr.").replace(" dr★", " dr.").replace("Dr★", "Dr.")
                f.write(f"{entry['number']}\n{start_time} --> {end_time}\n{text}\n\n")

        with open(srt_output_path, 'r', encoding='utf-8') as f:
            srt_content = f.read()

        txt_nr_content = ""
        for word_info in words_data:
            if not txt_nr_content:
                txt_nr_content += word_info['word'].lstrip()
            else:
                txt_nr_content += word_info['word']

        txt_nr_output_file_name = f"{input_file_name}_NR.txt"
        txt_nr_output_path = os.path.join(temp_dir, txt_nr_output_file_name)
        with open(txt_nr_output_path, 'w', encoding='utf-8') as f:
            txt_nr_content = txt_nr_content.replace(" Dr★", " Dr.").replace(" dr★", " dr.").replace("Dr★", "Dr.")
            f.write(txt_nr_content)

        txt_r_content = ""
        previous_word_end = 0
        is_first_word = True
        for word in words_data:
            if is_first_word or txt_r_content.endswith("\n"):
                txt_r_content += word['word'].strip()
            else:
                txt_r_content += word['word']
        
            if "." in word['word']:
                if word['start'] - previous_word_end >= 0.5:
                    txt_r_content += "\n"
                previous_word_end = word['end']
            is_first_word = False

        txt_r_output_file_name = f"{input_file_name}_R.txt"
        txt_r_output_path = os.path.join(temp_dir, txt_r_output_file_name)

        with open(txt_r_output_path, 'w', encoding='utf-8') as f:
            txt_r_content = txt_r_content.replace(" Dr★", " Dr.").replace(" dr★", " dr.").replace("Dr★", "Dr.")
            f.write(txt_r_content)

        # srtファイルからワードファイルへ変換
        doc_srt = Document()

        srtdoc_output_file_name = f"{input_file_name}_srt.docx"
        srtdoc_output_path = os.path.join(temp_dir, srtdoc_output_file_name)

        with open(srt_output_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()

        subtitle_number = None
        timestamp = None
        subtitle_text = []

        for line in lines:
            line = line.strip()
            if line.isdigit():
                # 以前の字幕エントリを追加
                if subtitle_number is not None and subtitle_text:
                    doc_srt.add_paragraph(f'{subtitle_number}')
                    doc_srt.add_paragraph(f'{timestamp}')
                    doc_srt.add_paragraph(' '.join(subtitle_text))
                    doc_srt.add_paragraph()  # 空行で区切る

                subtitle_number = line
                timestamp = None
                subtitle_text = []
            elif '-->' in line:
                timestamp = line
            elif line:
                subtitle_text.append(line)

        if subtitle_number is not None and subtitle_text:
            doc_srt.add_paragraph(f'{subtitle_number}')
            doc_srt.add_paragraph(f'{timestamp}')
            doc_srt.add_paragraph(' '.join(subtitle_text))

        doc_srt.save(srtdoc_output_path)

        ## txt(nr)をdoc変換
        txtdoc_nr = Document()
        txtdoc_nr_output_file_name = f"{input_file_name}_txtnr.docx"
        txtdoc_nr_output_path = os.path.join(temp_dir, txtdoc_nr_output_file_name)

        with open(txt_nr_output_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()

        for line in lines:
            txtdoc_nr.add_paragraph(line)

        txtdoc_nr.save(txtdoc_nr_output_path)

        ## txt(r)をdoc変換
        txtdoc_r = Document()
        txtdoc_r_output_file_name = f"{input_file_name}_txtr.docx"
        txtdoc_r_output_path = os.path.join(temp_dir, txtdoc_r_output_file_name)

        with open(txt_r_output_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()

        for line in lines:
            txtdoc_r.add_paragraph(line)

        txtdoc_r.save(txtdoc_r_output_path)

        #xls,df追加

        excel_filepath, df_display = create_excel_from_srt_c(srt_content=srt_content, input_file_name=input_file_name)

        # zipファイルにまとめる(srt,txtr,txtnr)。
        zip_core_file_name = f"{input_file_name}_core.zip"
        zip_core_file_path = os.path.join(temp_dir, zip_core_file_name)

        with zipfile.ZipFile(zip_core_file_path, 'w') as zip_file:
            zip_file.write(json_output_path, os.path.basename(json_output_path))
            zip_file.write(srt_output_path, os.path.basename(srt_output_path))
            zip_file.write(txt_r_output_path, os.path.basename(txt_r_output_path))
            zip_file.write(txt_nr_output_path, os.path.basename(txt_nr_output_path))
            
     
        # zipファイルにまとめる(doc)。
        zip_doc_file_name = f"{input_file_name}_office_en.zip"
        zip_doc_file_path = os.path.join(temp_dir, zip_doc_file_name)

        with zipfile.ZipFile(zip_doc_file_path, 'w') as zip_file:
            zip_file.write(excel_filepath, os.path.basename(excel_filepath))
            zip_file.write(txtdoc_nr_output_path, os.path.basename(txtdoc_nr_output_path))
            zip_file.write(txtdoc_r_output_path, os.path.basename(txtdoc_r_output_path))

        print(f"Processed {FileName}")
        
        main_files = [
            json_output_path,
            srt_output_path,
            txt_nr_output_path,
            txt_r_output_path,
            zip_core_file_path
        ]
        
        zip_doc_file_path = os.path.join(temp_dir, zip_doc_file_name)

        doc_files = [excel_filepath, txtdoc_nr_output_path, txtdoc_r_output_path, zip_doc_file_path]

        ##テーブル##

        df_display=dataframe_to_html_table(df_display)
        df_display=f"""
            <div class="my-table-container">
                {df_display}
            </div>
        """

        html_srt = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{srt_content}</pre>"""
        html_nr_txt = f"""<pre style="white-space: pre-wrap; overflow-y:auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{txt_nr_content}</pre>"""
        html_r_txt = f"""<pre style="white-space: pre-wrap; overflow-y:auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{txt_r_content}</pre>"""

        filename_copy = input_file_name
        srt_dummy_output_path = srt_output_path

        return srt_content, txt_nr_content, txt_r_content, main_files, doc_files ,html_srt, html_nr_txt, html_r_txt, filename_copy, srt_dummy_output_path, df_display
    except Exception as e:
        print(f"ファイル処理中にエラーが発生しました: {e}")
        traceback.print_exc()
        return"", "", "", [], [], "", "", "", "", "", ""

'''
テキストエリア①
テキストエリア②
テキストエリア③
ファイル④
ファイル⑤
HTML⑥
HTML⑦
HTML⑧
テキストボックス⑨
テキストボックス⑩
HTML⑪
'''
